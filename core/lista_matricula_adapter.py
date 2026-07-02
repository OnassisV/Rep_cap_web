"""Adaptador para el submenu "Lista de Matricula" (Operaciones de Plataforma).

Puerto de app_difoca/herramientas/list_matricula.py (Streamlit + SQLAlchemy)
a Django + pymysql, consultando oferta_formativa_difoca y bbdd_difoca en
Railway via accounts.db.get_connection(), siguiendo el patron ya usado en
core/legacy_adapters.py (obtener_cursos_aula_virtual, obtener_caracterizacion_por_dnis).
"""

import io
import logging
import zipfile
from datetime import datetime, date
from typing import Any

from accounts.db import get_connection

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt, RGBColor

logger = logging.getLogger(__name__)

ANIO_VIGENTE = datetime.now().year

TIPOS_LISTADO = ("Matriculados", "Participantes", "Certificados")

_COLUMNAS_VACIAS = {"", "-", "NO DISPONIBLE"}


def obtener_opciones_filtro_lista_matricula() -> dict[str, Any]:
    """Distinct anios/condiciones/procesos formativos para poblar los filtros.

    Equivalente a la carga de `oferta_inicial` del script legado.
    """
    query = """
        SELECT DISTINCT
            anio,
            condicion,
            tipo_proceso_formativo,
            denominacion_proceso_formativo
        FROM oferta_formativa_difoca
        WHERE anio IS NOT NULL
    """
    try:
        with get_connection() as connection:
            with connection.cursor() as cursor:
                cursor.execute(query)
                rows = list(cursor.fetchall())
    except Exception:
        logger.exception("Error en obtener_opciones_filtro_lista_matricula")
        rows = []

    anios = sorted({int(r["anio"]) for r in rows if r.get("anio")}, reverse=True)
    condiciones = sorted({str(r["condicion"]).strip() for r in rows if r.get("condicion")})

    procesos: list[dict[str, Any]] = []
    vistos: set[tuple] = set()
    for r in rows:
        anio = r.get("anio")
        condicion = str(r.get("condicion") or "").strip()
        proceso = f"{r.get('tipo_proceso_formativo', '')} {r.get('denominacion_proceso_formativo', '')}".strip()
        if not proceso:
            continue
        clave = (anio, condicion, proceso)
        if clave in vistos:
            continue
        vistos.add(clave)
        procesos.append({"anio": anio, "condicion": condicion, "proceso_formativo": proceso})

    procesos.sort(key=lambda p: (-(p["anio"] or 0), p["proceso_formativo"]))

    return {
        "anios": anios,
        "condiciones": condiciones,
        "procesos": procesos,
        "anio_vigente": ANIO_VIGENTE,
    }


def _parsear_fecha(valor: str | None) -> date | None:
    if not valor:
        return None
    try:
        return datetime.strptime(str(valor).strip(), "%Y-%m-%d").date()
    except ValueError:
        return None


def obtener_lista_matricula(
    anios: list[int] | None,
    condiciones: list[str] | None,
    procesos_formativos: list[str] | None,
    tipo_listado: str = "Matriculados",
    fecha_inicio: str | None = None,
    fecha_fin: str | None = None,
) -> dict[str, Any]:
    """Puerto de `obtener_lista_matricula` (list_matricula.py) a pymysql puro.

    Devuelve un dict con:
      - filas: list[dict] (N, REGION, NOMBRE_IGED, APELLIDOS, NOMBRES,
        NOMBRE_DE_PUESTO, [SITUACION_DEL_PARTICIPANTE], PROCESO_FORMATIVO,
        codigo, anio), ordenadas por REGION, NOMBRE_IGED, APELLIDOS, NOMBRES.
      - agregar_situacion: bool
      - regiones: list[str] (ordenadas, unicas, no vacias)
      - total, total_regiones, total_igeds, total_procesos: int
    """
    vacio = {
        "filas": [],
        "agregar_situacion": False,
        "regiones": [],
        "total": 0,
        "total_regiones": 0,
        "total_igeds": 0,
        "total_procesos": 0,
    }

    where = ["1=1"]
    params: list[Any] = []
    if anios:
        where.append("anio IN ({})".format(", ".join(["%s"] * len(anios))))
        params.extend(anios)
    if condiciones:
        where.append("condicion IN ({})".format(", ".join(["%s"] * len(condiciones))))
        params.extend(condiciones)

    oferta_query = f"""
        SELECT
            codigo,
            anio,
            condicion,
            tipo_proceso_formativo,
            denominacion_proceso_formativo,
            implementacion_final
        FROM oferta_formativa_difoca
        WHERE {' AND '.join(where)}
    """

    try:
        with get_connection() as connection:
            with connection.cursor() as cursor:
                cursor.execute(oferta_query, params)
                oferta_rows = list(cursor.fetchall())
    except Exception:
        logger.exception("Error consultando oferta_formativa_difoca en lista de matricula")
        return vacio

    # Filtro de fecha sobre Implementacion Final (se aplica en Python: el
    # rango es opcional y el volumen de filas de oferta es pequeño).
    fecha_inicio_dt = _parsear_fecha(fecha_inicio)
    fecha_fin_dt = _parsear_fecha(fecha_fin)
    if fecha_inicio_dt and fecha_fin_dt:
        filtradas = []
        for row in oferta_rows:
            valor = row.get("implementacion_final")
            fecha_val = valor.date() if isinstance(valor, datetime) else valor
            if isinstance(fecha_val, date) and fecha_inicio_dt <= fecha_val <= fecha_fin_dt:
                filtradas.append(row)
        oferta_rows = filtradas

    for row in oferta_rows:
        row["proceso_formativo"] = (
            f"{row.get('tipo_proceso_formativo', '')} {row.get('denominacion_proceso_formativo', '')}".strip()
        )

    if procesos_formativos:
        procesos_set = set(procesos_formativos)
        oferta_rows = [r for r in oferta_rows if r["proceso_formativo"] in procesos_set]

    codigos_oferta = sorted({
        str(r["codigo"]).strip()
        for r in oferta_rows
        if str(r.get("codigo") or "").strip()
    })
    if not codigos_oferta:
        return vacio

    bbdd_query = """
        SELECT
            codigo,
            estado,
            compromiso,
            aprobados_certificados,
            region,
            nombre_iged,
            apellidos,
            nombres,
            nombre_puesto,
            situacion_participante
        FROM bbdd_difoca
        WHERE codigo IN ({})
    """.format(", ".join(["%s"] * len(codigos_oferta)))

    try:
        with get_connection() as connection:
            with connection.cursor() as cursor:
                cursor.execute(bbdd_query, codigos_oferta)
                bbdd_rows = list(cursor.fetchall())
    except Exception:
        logger.exception("Error consultando bbdd_difoca en lista de matricula")
        return vacio

    bbdd_por_codigo: dict[str, list[dict[str, Any]]] = {}
    for row in bbdd_rows:
        bbdd_por_codigo.setdefault(str(row.get("codigo") or "").strip(), []).append(row)

    # tipo_listado determina el predicado sobre bbdd_difoca (estado/compromiso/certificados).
    def cumple_tipo_listado(fila: dict[str, Any]) -> bool:
        if fila.get("estado") != 2:
            return False
        if tipo_listado in ("Participantes", "Certificados") and fila.get("compromiso") != 20:
            return False
        if tipo_listado == "Certificados" and fila.get("aprobados_certificados") != 1:
            return False
        return True

    # Nota: replica una particularidad del script original — solo la
    # primera condicion seleccionada habilita la columna de situacion.
    primera_condicion = condiciones[0] if condiciones else None
    agregar_situacion = tipo_listado == "Participantes" and primera_condicion == "Cerrado"

    filas: list[dict[str, Any]] = []
    for oferta_row in oferta_rows:
        codigo = str(oferta_row.get("codigo") or "").strip()
        for bbdd_row in bbdd_por_codigo.get(codigo, []):
            if not cumple_tipo_listado(bbdd_row):
                continue

            nombre_iged = str(bbdd_row.get("nombre_iged") or "").strip().upper()

            nombre_puesto = str(bbdd_row.get("nombre_puesto") or "").strip().upper()
            if nombre_puesto in _COLUMNAS_VACIAS:
                nombre_puesto = "SIN REGISTRAR"

            fila = {
                "REGION": bbdd_row.get("region") or "",
                "NOMBRE_IGED": nombre_iged,
                "APELLIDOS": bbdd_row.get("apellidos") or "",
                "NOMBRES": bbdd_row.get("nombres") or "",
                "NOMBRE_DE_PUESTO": nombre_puesto,
                "PROCESO_FORMATIVO": oferta_row["proceso_formativo"],
                "codigo": codigo,
                "anio": oferta_row.get("anio"),
            }
            if agregar_situacion:
                situacion = str(bbdd_row.get("situacion_participante") or "").strip().lower()
                if situacion == "aprueba":
                    fila["SITUACION_DEL_PARTICIPANTE"] = "CERTIFICA"
                elif situacion == "no aprueba":
                    fila["SITUACION_DEL_PARTICIPANTE"] = "NO CERTIFICA"
                else:
                    fila["SITUACION_DEL_PARTICIPANTE"] = ""
            filas.append(fila)

    filas.sort(key=lambda f: (f["REGION"], f["NOMBRE_IGED"], f["APELLIDOS"], f["NOMBRES"]))
    for idx, fila in enumerate(filas, start=1):
        fila["N"] = idx

    regiones = sorted({f["REGION"] for f in filas if f["REGION"]})

    return {
        "filas": filas,
        "agregar_situacion": agregar_situacion,
        "regiones": regiones,
        "total": len(filas),
        "total_regiones": len(regiones),
        "total_igeds": len({f["NOMBRE_IGED"] for f in filas if f["NOMBRE_IGED"]}),
        "total_procesos": len({f["PROCESO_FORMATIVO"] for f in filas if f["PROCESO_FORMATIVO"]}),
    }


def obtener_resumen_por_region(filas: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Puerto de la agregacion `groupby('REGION')` (pestana Resumen, sin grafico)."""
    acumulado: dict[str, dict[str, Any]] = {}
    for fila in filas:
        region = fila.get("REGION") or ""
        if not region:
            continue
        entry = acumulado.setdefault(region, {"total_matriculados": 0, "igeds": set(), "procesos": set()})
        entry["total_matriculados"] += 1
        if fila.get("NOMBRE_IGED"):
            entry["igeds"].add(fila["NOMBRE_IGED"])
        if fila.get("PROCESO_FORMATIVO"):
            entry["procesos"].add(fila["PROCESO_FORMATIVO"])

    resumen = [
        {
            "region": region,
            "total_matriculados": data["total_matriculados"],
            "total_igeds": len(data["igeds"]),
            "total_procesos": len(data["procesos"]),
        }
        for region, data in acumulado.items()
    ]
    resumen.sort(key=lambda r: r["region"])
    return resumen


# ---------------------------------------------------------------------------
# Generacion de documentos Word
# ---------------------------------------------------------------------------

_COL_WIDTHS = [Cm(1.2), Cm(3.2), Cm(6.0), Cm(5.0), Cm(5.0), Cm(6.0)]
_COL_WIDTH_SITUACION = Cm(3.0)


def crear_documento_word(
    filas_region: list[dict[str, Any]],
    region: str,
    agregar_situacion: bool,
) -> Document:
    """Puerto de `crear_documento_word` (list_matricula.py); documento en memoria."""
    doc = Document()

    section = doc.sections[0]
    section.page_height = Cm(21.0)
    section.page_width = Cm(29.7)
    section.orientation = 1  # landscape

    procesos_en_region = sorted({f["PROCESO_FORMATIVO"] for f in filas_region})
    cols = 7 if agregar_situacion else 6

    if not procesos_en_region:
        doc.add_paragraph("No hay matriculados registrados para esta región.")
        return doc

    p_anexo = doc.add_paragraph()
    run_anexo = p_anexo.add_run("ANEXO")
    run_anexo.font.size = Pt(12)
    run_anexo.font.name = "Arial"
    run_anexo.bold = True
    p_anexo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for idx, proceso in enumerate(procesos_en_region):
        filas_proceso = sorted(
            (f for f in filas_region if f["PROCESO_FORMATIVO"] == proceso),
            key=lambda f: (f["NOMBRE_IGED"], f["APELLIDOS"], f["NOMBRES"]),
        )

        table = doc.add_table(rows=2, cols=cols)
        table.style = "Table Grid"

        # Fila 0: titulo del proceso formativo (celda combinada).
        title_cell = table.rows[0].cells[0]
        proceso_upper = str(proceso).upper()
        title_cell.text = proceso_upper
        for i in range(1, cols):
            table.rows[0].cells[i]._tc.merge(table.rows[0].cells[0]._tc)
        p_title = title_cell.paragraphs[0]
        run_title = p_title.runs[0] if p_title.runs else p_title.add_run(proceso_upper)
        run_title.font.size = Pt(11)
        run_title.font.name = "Arial"
        run_title.bold = True
        run_title.font.color.rgb = RGBColor(255, 255, 255)
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="808080"/>'.format(nsdecls("w"))))

        # Fila 1: encabezados.
        hdr_cells = table.rows[1].cells
        headers = ["N°", "REGIÓN", "NOMBRE IGED", "APELLIDOS", "NOMBRES", "NOMBRE DE PUESTO"]
        if agregar_situacion:
            headers.append("SITUACION DEL PARTICIPANTE")
        for i, header in enumerate(headers):
            p = hdr_cells[i].paragraphs[0]
            run = p.add_run(header)
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = "Arial"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells[i].vertical_alignment = 1
            hdr_cells[i]._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls("w"))))
            hdr_cells[i].width = _COL_WIDTHS[i] if i < 6 else _COL_WIDTH_SITUACION

        # Encabezado repetible en salto de pagina.
        for row_idx in (0, 1):
            tr = table.rows[row_idx]._tr
            trPr = tr.get_or_add_trPr()
            tblHeader = OxmlElement("w:tblHeader")
            tblHeader.set(qn("w:val"), "1")
            trPr.append(tblHeader)

        # Filas de datos.
        for numero, fila in enumerate(filas_proceso, start=1):
            row_cells = table.add_row().cells
            valores = [
                str(numero),
                str(fila.get("REGION", "")),
                str(fila.get("NOMBRE_IGED", "")).replace("\n", " ").replace("\r", " "),
                str(fila.get("APELLIDOS", "")).replace("\n", " ").replace("\r", " "),
                str(fila.get("NOMBRES", "")).replace("\n", " ").replace("\r", " "),
                str(fila.get("NOMBRE_DE_PUESTO", "")).replace("\n", " ").replace("\r", " "),
            ]
            if agregar_situacion:
                valores.append(str(fila.get("SITUACION_DEL_PARTICIPANTE", "")))
            for i, valor in enumerate(valores):
                cell = row_cells[i]
                cell.text = valor
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.name = "Arial"
                cell.width = _COL_WIDTHS[i] if i < 6 else _COL_WIDTH_SITUACION
                cell.vertical_alignment = 1

        for row in table.rows:
            row.height = Cm(1.0)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            for i, cell in enumerate(row.cells):
                cell.width = _COL_WIDTHS[i] if i < 6 else _COL_WIDTH_SITUACION
                cell.vertical_alignment = 1

        if idx < len(procesos_en_region) - 1:
            doc.add_paragraph("")

    return doc


def documento_word_a_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def crear_zip_regiones(
    filas: list[dict[str, Any]],
    regiones: list[str],
    agregar_situacion: bool,
    tipo_listado: str,
) -> bytes:
    """Construye un .zip en memoria con un .docx por region no vacia."""
    tipo_fmt = tipo_listado.replace(" ", "_").upper()
    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for region in regiones:
            filas_region = [f for f in filas if f.get("REGION") == region]
            if not filas_region:
                continue
            doc = crear_documento_word(filas_region, region, agregar_situacion)
            nombre = f"Lista_{tipo_fmt}_{region.replace(' ', '_')}.docx"
            zf.writestr(nombre, documento_word_a_bytes(doc))
    zip_buf.seek(0)
    return zip_buf.getvalue()
